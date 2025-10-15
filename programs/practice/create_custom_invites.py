# Chapter 17 - PDF and Word Documents
# (Practice Program) - Custom Invitations

import sys
import docx
import argparse

from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt


class ArgumentParserWithHelp(argparse.ArgumentParser):
    def error(self, message):
        """Override default error handling to show help text on invalid args."""
        self.print_help(sys.stderr)
        self.exit(2, f"\nError: {message}\n\n")


def get_lines(fp):
    doc = docx.Document(fp)
    lines = []
    for para in doc.paragraphs:
        lines.append(para.text)
    return lines


def get_guests(guest_file: str):
    try:
        with open(guest_file, 'r') as file:
            lines = file.readlines()
        return lines
    except FileNotFoundError:
        print(f"Guest file {guest_file} not found")
        sys.exit(1)


def create_custom_invites(guests_file: str, docx_file: str):
    guests = get_guests(guests_file)
    doc = docx.Document()
    lines = get_lines(docx_file)

    def style(p, line):
        r = p.add_run(line)
        r.bold = True
        r.font.size = Pt(18)
        r.italic = True
        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        return p

    for g in guests:
        style(doc.add_paragraph(), lines[0])
        p = style(doc.add_paragraph(), g)
        p.runs[0].italic = False
        for line in lines[1:]:
            style(doc.add_paragraph(), line)

        doc.add_page_break()
    
    doc.save("custom_invitations.docx")
    

def main():
    parser = ArgumentParserWithHelp(
        description="Create custom invites given a list of guests a invite template docx."
    )

    subparsers = parser.add_subparsers(dest="command", required=True)

    # --- send subcommand ---
    read_parser = subparsers.add_parser("read", help="Read a docx file")
    read_parser.add_argument("docx_file", type=str, help="Filepath of docx file")

    # --- listen subcommand ---
    write_parser = subparsers.add_parser("write", help="Create custom invites given a guest list txt file and a docx invite template")
    write_parser.add_argument("guest_list_file", type=str, help="Filepath of list of guests to create invitations for")
    write_parser.add_argument("docx_file", type=str, help="Filepath of invite template docx file")

    args = parser.parse_args()

    # --- handle subcommands ---
    if args.command == "read":
        try:
            if str(args.docx_file).endswith('.docx'):
                lines = get_lines(fp=args.docx_file)
                print(lines)
        except FileNotFoundError as fnfe:
            print(fnfe)
            sys.exit(1)
            
    elif args.command == "write":
        create_custom_invites(args.guest_list_file, args.docx_file)


if __name__ == "__main__":
    main()